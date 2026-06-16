"""Build the micro-expression article draft in LaTeX and DOCX formats."""

from __future__ import annotations

import csv
import shutil
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "article_draft"
FIGURES_DIR = ARTICLE_DIR / "figures"

RESULTS_DIR = ROOT / "microexpressions_repo" / "results"
BUNDLE_DIR = Path(r"C:\Users\gitya\Downloads\micro_exp_analysis_bundle")
ANNOTATIONS_CSV = BUNDLE_DIR / "inputs" / "annotations.csv"
WORD_SOURCE = Path(r"C:\Users\gitya\Downloads\הקובץ וורד.docx")
PPTX_SOURCE = Path(r"C:\Users\gitya\Downloads\התקדמות הפרויקט שלנו 5.pptx")

FIGURE_SOURCES = [
    RESULTS_DIR / "statistical_figures" / "quality_face_detection_rate.png",
    RESULTS_DIR / "statistical_figures" / "event_count_heatmap.png",
    RESULTS_DIR / "statistical_figures" / "event_duration_histogram.png",
    RESULTS_DIR / "statistical_figures" / "metric_correlation_heatmap.png",
    RESULTS_DIR / "statistical_figures" / "pca_cluster_plot.png",
]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def extract_docx_paragraphs(path: Path) -> list[str]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as package:
        root = ET.fromstring(package.read("word/document.xml"))
    paragraphs: list[str] = []
    for paragraph in root.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in paragraph.findall(".//w:t", ns)).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def collect_metrics() -> dict[str, object]:
    manifest = read_csv(RESULTS_DIR / "frame_features_manifest.csv")
    quality = read_csv(RESULTS_DIR / "quality_full.csv")
    events = read_csv(RESULTS_DIR / "auto_candidate_events.csv")
    sniffs = read_csv(RESULTS_DIR / "auto_candidate_sniffs.csv")
    annotations = read_csv(ANNOTATIONS_CSV)
    condition_mapping = read_csv(RESULTS_DIR / "condition_mapping.csv") if (RESULTS_DIR / "condition_mapping.csv").exists() else []
    condition_comparison = read_csv(RESULTS_DIR / "condition_comparison.csv") if (RESULTS_DIR / "condition_comparison.csv").exists() else []

    quality_counts = Counter(row["quality_tier"] for row in quality)
    annotation_flags = Counter(row["quality_flag"] or "blank" for row in annotations)
    metric_counts = Counter(row["metric"] for row in events)

    processed_frames = sum(int(float(row["processed_frames"])) for row in manifest)
    mean_face_detection = sum(float(row["face_detection_rate"]) for row in quality) / len(quality)
    event_durations = [float(row["duration_s"]) for row in events]
    event_peaks = [abs(float(row["peak_abs_z"])) for row in events]

    def median(values: list[float]) -> float:
        ordered = sorted(values)
        midpoint = len(ordered) // 2
        if len(ordered) % 2:
            return ordered[midpoint]
        return (ordered[midpoint - 1] + ordered[midpoint]) / 2

    annotation_videos = sorted({row["video_id"] for row in annotations})
    annotation_rows_by_video = {
        video_id: [row for row in annotations if row["video_id"] == video_id] for video_id in annotation_videos
    }
    four_trial_videos = sorted(
        video_id
        for video_id in annotation_videos
        if len(annotation_rows_by_video[video_id]) == 4
    )
    two_trial_videos = sorted(
        video_id
        for video_id in annotation_videos
        if len(annotation_rows_by_video[video_id]) == 2
    )
    two_trial_rows = [row for video_id in two_trial_videos for row in annotation_rows_by_video[video_id]]
    primary_two_trial_videos = sorted(
        video_id
        for video_id in two_trial_videos
        if all(
            (row["quality_flag"].strip().lower() == "ok") and bool(row["sniff_onset_s"].strip())
            for row in annotation_rows_by_video[video_id]
        )
    )
    primary_two_trial_rows = [
        row for video_id in primary_two_trial_videos for row in annotation_rows_by_video[video_id]
    ]
    pending_or_excluded_two_trial_rows = [
        row
        for row in two_trial_rows
        if row["video_id"] not in primary_two_trial_videos
        or row["quality_flag"].strip().lower() != "ok"
        or not row["sniff_onset_s"].strip()
    ]

    def to_float(value: str) -> float:
        try:
            return float(value)
        except (TypeError, ValueError):
            return float("nan")

    condition_pairs = [
        int(to_float(row["n_pairs"]))
        for row in condition_comparison
        if row.get("n_pairs") and to_float(row["n_pairs"]) == to_float(row["n_pairs"])
    ]
    top_condition_rows = sorted(
        condition_comparison,
        key=lambda row: to_float(row.get("paired_t_p", "")),
    )[:5]

    return {
        "manifest_rows": len(manifest),
        "processed_frames": processed_frames,
        "processed_videos": sum(1 for row in manifest if row["status"] in {"processed", "skipped_existing"}),
        "deepface_files": sum(1 for row in manifest if row["deepface_path"].strip()),
        "source_fps_values": sorted({row["source_fps"] for row in manifest if row["source_fps"]}),
        "quality_rows": len(quality),
        "quality_counts": quality_counts,
        "mean_face_detection": mean_face_detection,
        "events": len(events),
        "sniffs": len(sniffs),
        "event_metric_counts": metric_counts,
        "median_event_duration": median(event_durations),
        "median_peak_z": median(event_peaks),
        "annotations_rows": len(annotations),
        "annotation_videos": len(annotation_videos),
        "annotations_sniff_filled": sum(bool(row["sniff_onset_s"].strip()) for row in annotations),
        "annotations_approach_filled": sum(bool(row["approach_onset_s"].strip()) for row in annotations),
        "annotations_response_filled": sum(bool(row["response_s"].strip()) for row in annotations),
        "annotation_flags": annotation_flags,
        "four_trial_videos": four_trial_videos,
        "two_trial_videos": two_trial_videos,
        "two_trial_rows": len(two_trial_rows),
        "primary_two_trial_videos": primary_two_trial_videos,
        "primary_two_trial_rows": len(primary_two_trial_rows),
        "pending_or_excluded_two_trial_rows": len(pending_or_excluded_two_trial_rows),
        "condition_mapping_rows": len(condition_mapping),
        "condition_mapping_videos": len({row["video_id"] for row in condition_mapping}) if condition_mapping else 0,
        "condition_comparison_rows": len(condition_comparison),
        "condition_n_pairs": max(condition_pairs) if condition_pairs else 0,
        "top_condition_rows": top_condition_rows,
    }


def latex_escape(value: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in value)


def write_references() -> None:
    references = r"""@inproceedings{baltrusaitis2018openface,
  title={OpenFace 2.0: Facial Behavior Analysis Toolkit},
  author={Baltru{\v{s}}aitis, Tadas and Zadeh, Amir and Lim, Yao Chong and Morency, Louis-Philippe},
  booktitle={2018 13th IEEE International Conference on Automatic Face \& Gesture Recognition},
  pages={59--66},
  year={2018},
  publisher={IEEE}
}

@incollection{cohn2015automated,
  title={Automated Face Analysis for Affective Computing},
  author={Cohn, Jeffrey F. and De la Torre, Fernando},
  booktitle={The Oxford Handbook of Affective Computing},
  editor={Calvo, Rafael A. and D'Mello, Sidney and Gratch, Jonathan and Kappas, Arvid},
  year={2015},
  publisher={Oxford University Press}
}

@article{degroot2015sniff,
  title={A Sniff of Happiness},
  author={de Groot, Jasper H. B. and Smeets, Monique A. M. and Rowson, Michael J. and Bulsing, Peter J. and Blonk, Caroline G. and Wilkinson, John E. and Semin, G{\"u}n R.},
  journal={Psychological Science},
  volume={26},
  number={6},
  pages={684--700},
  year={2015}
}

@book{ekman1978facs,
  title={Facial Action Coding System: A Technique for the Measurement of Facial Movement},
  author={Ekman, Paul and Friesen, Wallace V.},
  year={1978},
  publisher={Consulting Psychologists Press}
}

@phdthesis{hatcher2016chemical,
  title={Chemical Communication: The Effects of Stress-Induced Apocrine Sweat on Human Perceptions and Interactions},
  author={Hatcher, Lauren},
  school={Louisiana State University},
  year={2016}
}

@article{hayduk1978personal,
  title={Personal Space: An Evaluative and Orienting Overview},
  author={Hayduk, Leslie A.},
  journal={Psychological Bulletin},
  volume={85},
  number={1},
  pages={117--134},
  year={1978}
}

@article{lebert2024distance,
  title={Keeping Distance or Getting Closer: How Others' Emotions Shape Approach-Avoidance Postural Behaviors and Preferred Interpersonal Distance},
  author={Lebert, Alexandra and Vergilino-Perez, Dorine and Chaby, Laurence},
  journal={PLOS ONE},
  volume={19},
  number={2},
  pages={e0298069},
  year={2024}
}

@article{lloyd2009space,
  title={The Space Between Us: A Neurophilosophical Framework for the Investigation of Human Interpersonal Space},
  author={Lloyd, Donna M.},
  journal={Neuroscience and Biobehavioral Reviews},
  volume={33},
  number={3},
  pages={297--304},
  year={2009}
}

@article{perry2015ot,
  title={Oxytocin Promotes Closer Interpersonal Distance Among Highly Empathic Individuals},
  author={Perry, Anat and Mankuta, David and Shamay-Tsoory, Simone G.},
  journal={Social Cognitive and Affective Neuroscience},
  volume={10},
  number={1},
  pages={3--9},
  year={2015}
}

@article{ruggiero2017facial,
  title={The Effect of Facial Expressions on Peripersonal and Interpersonal Spaces},
  author={Ruggiero, Gennaro and Frassinetti, Francesca and Coello, Yann and Rapuano, Maria and Di Cola, Angela S. and Iachini, Tina},
  journal={Psychological Research},
  volume={81},
  number={6},
  pages={1232--1240},
  year={2017}
}

@inproceedings{serengil2020lightface,
  title={LightFace: A Hybrid Deep Face Recognition Framework},
  author={Serengil, Sefik Ilkin and Ozpinar, Alper},
  booktitle={2020 Innovations in Intelligent Systems and Applications Conference},
  year={2020},
  publisher={IEEE}
}

@inproceedings{serengil2021hyperextended,
  title={HyperExtended LightFace: A Facial Attribute Analysis Framework},
  author={Serengil, Sefik Ilkin and Ozpinar, Alper},
  booktitle={2021 International Conference on Engineering and Emerging Technologies},
  year={2021},
  publisher={IEEE}
}

@article{tian2001recognizing,
  title={Recognizing Action Units for Facial Expression Analysis},
  author={Tian, Ying-li and Kanade, Takeo and Cohn, Jeffrey F.},
  journal={IEEE Transactions on Pattern Analysis and Machine Intelligence},
  volume={23},
  number={2},
  pages={97--115},
  year={2001}
}

@article{yan2013microexpression,
  title={How Fast Are the Leaked Facial Expressions: The Duration of Micro-Expressions},
  author={Yan, Wen-Jing and Wu, Qi and Liang, Jing and Chen, Yu-Hsin and Fu, Xiaolan},
  journal={Journal of Nonverbal Behavior},
  volume={37},
  pages={217--230},
  year={2013}
}

@misc{zagoorysharon2025poster,
  title={You Smell It, You Act on It, But You Don't Know It: The Unconscious Impact of Chemosignals on Interpersonal Distance},
  author={Zagoory-Sharon, Orna and Tonin Agranionih, M. and Shepherd, G.},
  note={Poster, Baruch Ivcher School of Psychology, Reichman University},
  year={2025}
}
"""
    (ARTICLE_DIR / "references.bib").write_text(references, encoding="utf-8")


def figure_filename(source: Path) -> str:
    return source.name


def copy_figures() -> None:
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    for source in FIGURE_SOURCES:
        if not source.exists():
            raise FileNotFoundError(source)
        shutil.copy2(source, FIGURES_DIR / source.name)


def article_sections(metrics: dict[str, object]) -> list[tuple[str, list[str]]]:
    quality_counts: Counter[str] = metrics["quality_counts"]  # type: ignore[assignment]
    annotation_flags: Counter[str] = metrics["annotation_flags"]  # type: ignore[assignment]
    metric_counts: Counter[str] = metrics["event_metric_counts"]  # type: ignore[assignment]
    metric_labels = {
        "au45_blink_proxy": "AU45 blink",
        "au06_cheek_raise_proxy": "AU6 cheek raise",
        "au09_nose_wrinkle_proxy": "AU9 nose wrinkle",
        "cheek_composite_proxy": "cheek composite",
        "au10_upper_lip_raise_proxy": "AU10 upper-lip raise",
        "au24_lip_press_proxy": "AU24 lip press",
        "mouth_open_proxy": "mouth opening",
        "head_yaw_proxy": "head yaw",
        "mouth_asymmetry_proxy": "mouth asymmetry",
        "au12_lip_corner_pull_proxy": "AU12 lip-corner pull",
    }
    top_metrics = ", ".join(
        f"{metric_labels.get(metric, metric)}={count}" for metric, count in metric_counts.most_common(6)
    )
    suffixes = ("_z_peak_abs", "_delta_mean", "_delta_auc_abs", "_raw_mean")

    def clean_feature_label(feature: str) -> str:
        label = feature
        for suffix in suffixes:
            label = label.replace(suffix, "")
        return metric_labels.get(label, label.replace("_", " "))

    def safe_float(value: object) -> float:
        try:
            return float(value)  # type: ignore[arg-type]
        except (TypeError, ValueError):
            return float("nan")

    top_condition_rows: list[dict[str, str]] = metrics["top_condition_rows"]  # type: ignore[assignment]
    if top_condition_rows:
        formatted_condition_results = []
        for row in top_condition_rows[:3]:
            p_value = safe_float(row.get("paired_t_p", ""))
            diff = safe_float(row.get("mean_diff_sweat_minus_control", ""))
            dz = safe_float(row.get("paired_effect_dz", ""))
            formatted_condition_results.append(
                f"{row.get('window', 'window')} / {clean_feature_label(row.get('feature', 'feature'))}: "
                f"mean sweat-control difference={diff:.3g}, dz={dz:.2g}, paired t p={p_value:.3g}"
            )
        condition_result_sentence = " The lowest-p preliminary rows were " + "; ".join(formatted_condition_results) + "."
    else:
        condition_result_sentence = ""

    return [
        (
            "Abstract",
            [
                (
                    "Human sweat can contain chemosensory information that influences social perception and action even when people do not explicitly identify the source of the odor. "
                    "The present work develops a privacy-preserving computational pipeline for testing whether sweat-related chemosignals are accompanied by rapid facial responses and candidate micro-expressions during a filmed odor-discrimination task. "
                    f"The current dataset contains {metrics['manifest_rows']} representative videos processed at 100 frames per second, yielding {metrics['processed_frames']:,} frame-level observations, {metrics['events']:,} automatically detected candidate micro-events, and {metrics['sniffs']} candidate sniff-like windows. "
                    f"A newly added annotation file includes {metrics['annotations_rows']} trial rows across {metrics['annotation_videos']} videos, with {metrics['annotations_sniff_filled']} filled sniff onsets, {annotation_flags.get('ok', 0)} rows marked ok, {annotation_flags.get('pending', 0)} pending rows, and {annotation_flags.get('exclude', 0)} excluded rows. "
                    f"For the two-test-tube protocol, trial 1 was coded as placebo/control and trial 2 as sweat, producing {metrics['condition_mapping_rows']} mapped trial rows and a primary paired subset of {metrics['condition_n_pairs']} complete videos. "
                    "The resulting sweat-control analyses are reported as preliminary exploratory evidence rather than as a confirmatory test."
                )
            ],
        ),
        (
            "Introduction",
            [
                (
                    "Human body odor is not merely a nuisance variable in social interaction. A growing literature suggests that sweat and other chemosensory cues can carry affective and social information, sometimes influencing perception or behavior without being explicitly recognized by the perceiver. "
                    "Studies of chemosignal communication, interpersonal distance, and affective processing converge on the idea that olfactory cues may shape approach, avoidance, arousal, and evaluative response through pathways that are closely linked to limbic and embodied social systems."
                ),
                (
                    "The present project extends that rationale from overt interpersonal behavior to the face itself. In the two-test-tube paradigm, filmed participants smell two coded tubes, one intended to contain human sweat collected in a separate exercise-based sweat-donation study and one serving as a control stimulus. "
                    "The scientific question is whether the facial response contains an implicit temporal signature of sweat exposure, even when a participant's declarative identification of the tube is inaccurate or incomplete."
                ),
                (
                    "Micro-expressions and brief facial movements are theoretically relevant because they may reveal rapid affective or evaluative processing before the participant can regulate or verbalize the response. "
                    "The Facial Action Coding System decomposes facial behavior into action units, and automated facial-analysis tools make it possible to summarize movement patterns over time. "
                    "For odor exposure, the nose, upper lip, mouth, and cheek regions are especially important: AU9 and AU10 are associated with nose wrinkling and upper-lip raising, AU6 and AU12 capture cheek and lip-corner movement, and AU45 offers a blink-related marker that can change with sensory sampling or discomfort."
                ),
                (
                    "This article therefore describes both the planned experimental logic and the implemented computational video-analysis pipeline. It deliberately separates the methodological achievement of extracting high-frequency facial features from the still-unfinished experimental inference about sweat versus control. "
                    "That separation is essential: condition should never be inferred from file names or visual inspection. In the present revision, order is used only where the two-test-tube protocol was explicitly clarified: the first tube was placebo/control and the second tube was sweat."
                ),
            ],
        ),
        (
            "Research Questions and Hypotheses",
            [
                (
                    "The primary research question is whether dynamic facial-expression measures differ after sniffing a sweat tube compared with a control tube. "
                    "A second question asks whether this difference can be observed even when explicit tube identification is incorrect, which would support a partial or implicit chemosensory response. "
                    "A third question concerns which measures and time windows are most informative, with a particular emphasis on cheek, nose, upper-lip, blink, and mouth-related proxies."
                ),
                (
                    "The working hypotheses are directional but remain exploratory in the current version. Sweat trials are expected to show faster or stronger changes in measures linked to aversion, alertness, cautious sampling, or facial tension. "
                    "Early windows after sniff onset are expected to be more sensitive to micro-expression-like responses than later windows, whereas later windows may include more deliberate expression regulation. "
                    "The cheek composite is treated as a preregistered analytic priority because it summarizes AU6/AU12-related movement and adjacent geometric change."
                ),
            ],
        ),
        (
            "Method",
            [
                (
                    "The project distinguishes two samples. Sweat donors are the individuals from whom sweat samples were collected in a separate spinning-exercise study. Filmed participants are a separate group who complete the odor-exposure task. "
                    "This distinction matters ethically and analytically: the person providing the sweat sample is not the person whose face is analyzed in the video."
                ),
                (
                    "The target design is within-subjects. Each filmed participant smells coded tubes, and the central independent variable is stimulus type: sweat versus control. "
                    "The principal dependent variables are facial-expression proxies, candidate micro-events, latency to response, peak response, area under the curve, explicit identification accuracy, and trial-level quality measures. "
                    "For the two-test-tube protocol analyzed here, the first tube was coded as placebo/control and the second tube was coded as sweat. This rule was not applied to four-trial odor sequences."
                ),
                (
                    "The current video inventory contains 19 MP4 files. Four exact duplicates were identified and excluded from representative counts, leaving 15 representative videos for automatic analysis. "
                    f"The annotations contain {len(metrics['two_trial_videos'])} two-test-tube videos and {len(metrics['four_trial_videos'])} four-trial odor-sequence videos. "
                    "The two-test-tube videos define the sweat/control analysis set; four-trial rows are retained for traceability and exploratory timing work but are not merged into the sweat/control hypothesis test."
                ),
            ],
        ),
        (
            "Computational Video-Analysis Pipeline",
            [
                (
                    "The implemented pipeline reads source videos in place, leaving the original recordings unchanged. It uses OpenCV for video access, MediaPipe-derived facial landmarks for geometry-based measurement, and DeepFace as a low-frequency auxiliary emotion layer. "
                    "DeepFace is not treated as a substitute for FACS or for micro-expression measurement; it is used only as an additional descriptive layer."
                ),
                (
                    "The full representative dataset was processed at 100 frames per second. For each frame, the system extracted geometry-based proxies for cheek raising, lip-corner movement, nose wrinkling, upper-lip raising, lip press or mouth opening, blinking, asymmetry, and head movement. "
                    "The pipeline then summarized short-lived deviations from a rolling baseline as candidate micro-events when they exceeded a robust z-score threshold and lasted within the target micro-event duration range."
                ),
                (
                    "The current extraction produced "
                    f"{metrics['processed_frames']:,} processed frame rows across {metrics['manifest_rows']} representative videos. "
                    f"DeepFace sample files were present for {metrics['deepface_files']} videos. "
                    f"The automatic detector found {metrics['events']:,} candidate micro-events; the most frequent metrics were {top_metrics}. "
                    f"The median event duration was {metrics['median_event_duration']:.3f} s and the median peak absolute robust z-score was {metrics['median_peak_z']:.2f}."
                ),
            ],
        ),
        (
            "Annotation and Quality Control",
            [
                (
                    "Manual annotation remains the key step separating exploratory feature extraction from confirmatory experimental inference. "
                    f"The current annotations file contains {metrics['annotations_rows']} rows across {metrics['annotation_videos']} videos. "
                    f"Sniff onset is filled in {metrics['annotations_sniff_filled']} rows; approach onset and response time are not yet filled in the current annotations file. "
                    f"The quality flags are ok={annotation_flags.get('ok', 0)}, pending={annotation_flags.get('pending', 0)}, and exclude={annotation_flags.get('exclude', 0)}."
                ),
                (
                    f"The condition mapping now contains {metrics['condition_mapping_rows']} rows from the two-test-tube protocol. "
                    f"The primary paired analysis uses {metrics['condition_n_pairs']} complete videos ({metrics['primary_two_trial_rows']} trial rows) in which both trials were marked ok and had a sniff onset. "
                    f"{metrics['pending_or_excluded_two_trial_rows']} two-test-tube rows remain pending or excluded and were not included in the paired comparison."
                ),
                (
                    "Quality control was computed at the video level. "
                    f"The mean face-detection rate was {metrics['mean_face_detection']:.3f}. "
                    f"The quality tiers were good={quality_counts.get('good', 0)}, usable_review={quality_counts.get('usable_review', 0)}, and review_or_exclude={quality_counts.get('review_or_exclude', 0)}. "
                    "The lower-quality videos should be reviewed manually before any final group-level inference."
                ),
                (
                    "The analysis uses strict privacy boundaries. The article includes only aggregate plots and numeric summaries. It does not include participant frames, face crops, contact sheets, WhatsApp images, or other potentially identifying visual material."
                ),
            ],
        ),
        (
            "Preliminary Exploratory Results",
            [
                (
                    "The current results show that the computational layer is operational across the full representative dataset. The system successfully generated frame-level features, video-level quality summaries, candidate micro-events, candidate sniff-like windows, correlation plots, and exploratory PCA/clustering summaries. "
                    "These results are best understood as a feasibility and data-readiness outcome."
                ),
                (
                    f"The automatic sniff-window detector identified {metrics['sniffs']} candidate sniff-like windows. These windows are useful for prioritizing human review, but they are not final event labels. "
                    f"The trial-level paired report generated {metrics['condition_comparison_rows']} sweat-control comparison rows across baseline-normalized AU-proxy summaries. "
                    f"These comparisons use n={metrics['condition_n_pairs']} paired videos and should be interpreted as preliminary because the sample is small and the facial measures are automated proxies.{condition_result_sentence}"
                ),
                (
                    "The non-identifying figures included with this draft visualize the quality distribution, event density, event-duration distribution, correlations among facial metrics, and unsupervised PCA and clustering structure. "
                    "They support pipeline validation and exploratory profiling. The condition-level tables add a first paired sweat/control layer, but they do not by themselves establish a final condition-specific chemosignal effect."
                ),
            ],
        ),
        (
            "Discussion",
            [
                (
                    "The main contribution of the present work is a reproducible bridge between a psychologically motivated chemosignal question and a high-frequency computational facial-analysis workflow. "
                    "The pipeline turns raw video into interpretable numeric summaries while preserving participant privacy and making uncertainty visible through explicit quality-control fields."
                ),
                (
                    "The project is now positioned for a stronger next analysis step. The two-test-tube subset already has a protocol-based sweat/control mapping, allowing paired tests and Wilcoxon signed-rank tests on baseline-normalized trial summaries. "
                    "As more timings and condition metadata are completed, permutation tests and mixed-effects models will be appropriate for testing robustness and handling participant-level variability."
                ),
                (
                    "The current annotation file also clarifies an important design issue. Some rows appear to reflect a four-trial odor protocol rather than the planned two-test-tube sweat/control paradigm. "
                    "Rather than forcing these data into the main model, the article treats them as a separate protocol category requiring supervisor clarification. This conservative handling protects the validity of any later inference."
                ),
            ],
        ),
        (
            "Limitations and Next Steps",
            [
                (
                    "The central limitation is no longer the absence of condition labels for the two-test-tube subset, but the limited size and completeness of the analyzable paired sample. "
                    f"The current primary condition analysis is based on {metrics['condition_n_pairs']} complete paired videos, while pending/excluded two-test-tube rows and all four-trial rows remain outside the main sweat/control test. "
                    "A second limitation is that automatic facial geometry proxies are not FACS-certified AU labels; they should be treated as measurable approximations that require careful interpretation and, ideally, human-coded validation on a subset."
                ),
                (
                    "Immediate next steps are to confirm the remaining pending sniff onsets, decide whether the four-trial odor protocol has its own condition mapping, and review the videos flagged by quality control. "
                    "After manual review, the trial-level extraction and analysis commands should be rerun so that paired statistical tests and model-ready tables are based on validated experimental inputs."
                ),
            ],
        ),
    ]


def write_main_tex(metrics: dict[str, object]) -> None:
    sections = article_sections(metrics)
    figure_blocks = [
        (
            "quality_face_detection_rate.png",
            "Video-level face-detection quality across the representative dataset.",
            "fig:quality",
        ),
        (
            "event_count_heatmap.png",
            "Candidate micro-event density by video and metric.",
            "fig:event-heatmap",
        ),
        (
            "event_duration_histogram.png",
            "Distribution of automatically detected candidate event durations.",
            "fig:event-duration",
        ),
        (
            "metric_correlation_heatmap.png",
            "Correlations among video-level facial metric summaries.",
            "fig:correlations",
        ),
        (
            "pca_cluster_plot.png",
            "Exploratory PCA and clustering profile across representative videos.",
            "fig:pca",
        ),
    ]

    out: list[str] = [
        r"\documentclass[11pt]{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{graphicx}",
        r"\usepackage{booktabs}",
        r"\usepackage{float}",
        r"\usepackage{hyperref}",
        r"\usepackage[numbers]{natbib}",
        r"\usepackage{caption}",
        r"\usepackage{setspace}",
        r"\onehalfspacing",
        r"\title{Sweat Chemosignals and Facial Micro-Expressions:\\A Computational Video-Analysis Pipeline and Preliminary Exploratory Evidence}",
        r"\author{Yogev Git \and Ofir Goldreich \and Nir Shoham \\ Baruch Ivcher School of Psychology, Reichman University}",
        r"\date{Draft prepared June 2026}",
        r"\begin{document}",
        r"\maketitle",
        r"\noindent\textbf{Privacy note.} This draft includes only aggregate statistics and non-identifying plots. It does not include participant frames, face crops, contact sheets, or WhatsApp images.",
        "",
    ]

    for title, paragraphs in sections:
        command = "section*" if title == "Abstract" else "section"
        out.append(f"\\{command}{{{latex_escape(title)}}}")
        for paragraph in paragraphs:
            out.append(latex_escape(paragraph))
            out.append("")
        if title == "Preliminary Exploratory Results":
            out.append(r"\begin{table}[H]")
            out.append(r"\centering")
            out.append(r"\caption{Current data-readiness summary.}")
            out.append(r"\begin{tabular}{ll}")
            out.append(r"\toprule")
            out.append(r"Item & Current value \\")
            out.append(r"\midrule")
            out.append(f"Representative videos processed & {metrics['manifest_rows']}/15 \\\\")
            out.append(f"Processed frame rows & {metrics['processed_frames']:,} \\\\")
            out.append(f"Candidate micro-events & {metrics['events']:,} \\\\")
            out.append(f"Candidate sniff-like windows & {metrics['sniffs']} \\\\")
            out.append(f"Annotation rows & {metrics['annotations_rows']} \\\\")
            out.append(f"Filled sniff onsets & {metrics['annotations_sniff_filled']} \\\\")
            out.append(f"Two-test-tube condition mapping rows & {metrics['condition_mapping_rows']} \\\\")
            out.append(f"Primary paired videos & {metrics['condition_n_pairs']} \\\\")
            out.append(f"Paired sweat/control comparison rows & {metrics['condition_comparison_rows']} \\\\")
            out.append(r"\bottomrule")
            out.append(r"\end{tabular}")
            out.append(r"\end{table}")
            out.append("")

            for filename, caption, label in figure_blocks:
                out.append(r"\begin{figure}[H]")
                out.append(r"\centering")
                out.append(rf"\includegraphics[width=0.92\linewidth]{{figures/{filename}}}")
                out.append(rf"\caption{{{latex_escape(caption)}}}")
                out.append(rf"\label{{{label}}}")
                out.append(r"\end{figure}")
                out.append("")

    out.extend(
        [
            r"\nocite{*}",
            r"\bibliographystyle{plainnat}",
            r"\bibliography{references}",
            r"\end{document}",
        ]
    )

    (ARTICLE_DIR / "main.tex").write_text("\n".join(out), encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")
        borders.append(element)
    tbl_pr.append(borders)


def add_table(document: Document, rows: list[tuple[str, str]], caption: str) -> None:
    caption_p = document.add_paragraph()
    caption_run = caption_p.add_run(caption)
    caption_run.bold = True
    caption_run.font.size = Pt(10)
    caption_p.paragraph_format.space_before = Pt(8)
    caption_p.paragraph_format.space_after = Pt(4)

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.table_direction = WD_TABLE_DIRECTION.LTR
    table.columns[0].width = Inches(2.25)
    table.columns[1].width = Inches(4.0)
    set_table_borders(table)
    header = table.rows[0].cells
    header[0].text = "Item"
    header[1].text = "Current value"
    for cell in header:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for item, value in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = value
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_figure(document: Document, filename: str, caption: str) -> None:
    path = FIGURES_DIR / filename
    document.add_picture(str(path), width=Inches(6.0))
    last = document.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)


def write_docx(metrics: dict[str, object]) -> None:
    document = Document()
    configure_docx_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Sweat Chemosignals and Facial Micro-Expressions")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("A Computational Video-Analysis Pipeline and Preliminary Exploratory Evidence")
    subrun.italic = True
    subrun.font.size = Pt(13)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Yogev Git, Ofir Goldreich, and Nir Shoham\n").bold = True
    meta.add_run("Baruch Ivcher School of Psychology, Reichman University\nDraft prepared June 2026")

    note = document.add_paragraph()
    note_run = note.add_run("Privacy note. ")
    note_run.bold = True
    note.add_run(
        "This draft includes only aggregate statistics and non-identifying plots. It does not include participant frames, face crops, contact sheets, or WhatsApp images."
    )

    for title_text, paragraphs in article_sections(metrics):
        document.add_heading(title_text, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        if title_text == "Preliminary Exploratory Results":
            add_table(
                document,
                [
                    ("Representative videos processed", f"{metrics['manifest_rows']}/15"),
                    ("Processed frame rows", f"{metrics['processed_frames']:,}"),
                    ("Candidate micro-events", f"{metrics['events']:,}"),
                    ("Candidate sniff-like windows", str(metrics["sniffs"])),
                    ("Annotation rows", str(metrics["annotations_rows"])),
                    ("Filled sniff onsets", str(metrics["annotations_sniff_filled"])),
                    ("Two-test-tube condition mapping rows", str(metrics["condition_mapping_rows"])),
                    ("Primary paired videos", str(metrics["condition_n_pairs"])),
                    ("Paired sweat/control comparison rows", str(metrics["condition_comparison_rows"])),
                ],
                "Table 1. Current data-readiness summary.",
            )
            add_figure(
                document,
                "quality_face_detection_rate.png",
                "Figure 1. Video-level face-detection quality across the representative dataset.",
            )
            add_figure(
                document,
                "event_count_heatmap.png",
                "Figure 2. Candidate micro-event density by video and metric.",
            )
            add_figure(
                document,
                "event_duration_histogram.png",
                "Figure 3. Distribution of automatically detected candidate event durations.",
            )
            add_figure(
                document,
                "metric_correlation_heatmap.png",
                "Figure 4. Correlations among video-level facial metric summaries.",
            )
            add_figure(
                document,
                "pca_cluster_plot.png",
                "Figure 5. Exploratory PCA and clustering profile across representative videos.",
            )

    document.add_heading("References", level=1)
    refs = [
        "Baltrušaitis, T., Zadeh, A., Lim, Y. C., & Morency, L.-P. (2018). OpenFace 2.0: Facial behavior analysis toolkit. 13th IEEE International Conference on Automatic Face & Gesture Recognition, 59-66.",
        "Cohn, J. F., & De la Torre, F. (2015). Automated face analysis for affective computing. In R. A. Calvo et al. (Eds.), The Oxford Handbook of Affective Computing.",
        "De Groot, J. H. B., et al. (2015). A sniff of happiness. Psychological Science, 26(6), 684-700.",
        "Ekman, P., & Friesen, W. V. (1978). Facial Action Coding System. Consulting Psychologists Press.",
        "Hatcher, L. (2016). Chemical communication: The effects of stress-induced apocrine sweat. LSU Doctoral Dissertations.",
        "Hayduk, L. A. (1978). Personal space: An evaluative and orienting overview. Psychological Bulletin, 85(1), 117-134.",
        "Lebert, A., Vergilino-Perez, D., & Chaby, L. (2024). Keeping distance or getting closer. PLOS ONE, 19(2), e0298069.",
        "Lloyd, D. M. (2009). The space between us. Neuroscience and Biobehavioral Reviews, 33(3), 297-304.",
        "Perry, A., Mankuta, D., & Shamay-Tsoory, S. G. (2015). OT promotes closer interpersonal distance. Social Cognitive and Affective Neuroscience, 10(1), 3-9.",
        "Ruggiero, G., et al. (2017). The effect of facial expressions on peripersonal and interpersonal spaces. Psychological Research, 81(6), 1232-1240.",
        "Serengil, S. I., & Ozpinar, A. (2020, 2021). LightFace / HyperExtended LightFace.",
        "Tian, Y.-I., Kanade, T., & Cohn, J. F. (2001). Recognizing action units for facial expression analysis. IEEE Transactions on Pattern Analysis and Machine Intelligence, 23(2), 97-115.",
        "Yan, W.-J., Wu, Q., Liang, J., Chen, Y.-H., & Fu, X. (2013). How fast are the leaked facial expressions. Journal of Nonverbal Behavior, 37, 217-230.",
        "Zagoory-Sharon, O., Tonin Agranionih, M., & Shepherd, G. (2025). You smell it, you act on it, but you don't know it. Poster, Reichman University.",
    ]
    for ref in refs:
        paragraph = document.add_paragraph(ref)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.size = Pt(10)

    primary_path = ARTICLE_DIR / "microexpressions_article.docx"
    verified_path = ARTICLE_DIR / "microexpressions_article_verified.docx"
    document.save(verified_path)
    try:
        document.save(primary_path)
    except PermissionError:
        pass


def write_readme(metrics: dict[str, object]) -> None:
    text = f"""# Article Draft Sources

Generated deliverables:

- `main.tex`: LaTeX article source for Overleaf.
- `references.bib`: BibTeX references.
- `microexpressions_article_verified.docx`: Word version of the same article.
- `figures/`: non-identifying statistical figures only.

Primary sources used:

- Hebrew Word source: `C:/Users/gitya/Downloads/הקובץ וורד.docx`
- Project presentation: `C:/Users/gitya/Downloads/התקדמות הפרויקט שלנו 5.pptx`
- New annotations file: `C:/Users/gitya/Downloads/micro_exp_analysis_bundle/inputs/annotations.csv`
- Statistical outputs from `microexpressions_repo/results`

Verified numeric anchors:

- Representative videos processed: {metrics['manifest_rows']}/15
- Processed frame rows: {metrics['processed_frames']:,}
- Candidate micro-events: {metrics['events']:,}
- Candidate sniff-like windows: {metrics['sniffs']}
- Mean face detection rate: {metrics['mean_face_detection']:.3f}
- Quality tiers: {dict(metrics['quality_counts'])}
- Annotation rows: {metrics['annotations_rows']}
- Annotation videos: {metrics['annotation_videos']}
- Filled sniff onsets: {metrics['annotations_sniff_filled']}
- Annotation quality flags: {dict(metrics['annotation_flags'])}
- Two-test-tube videos: {len(metrics['two_trial_videos'])}
- Four-trial exploratory videos: {len(metrics['four_trial_videos'])}
- Condition mapping rows: {metrics['condition_mapping_rows']}
- Primary paired videos: {metrics['condition_n_pairs']}
- Paired sweat/control comparison rows: {metrics['condition_comparison_rows']}

Privacy and interpretation boundaries:

- Participant frames, face crops, contact sheets, and WhatsApp images were not copied or embedded.
- Figures are aggregate/statistical outputs only.
- Trial 1 in the two-test-tube protocol is coded as placebo/control; trial 2 is coded as sweat.
- The article treats sweat/control findings as preliminary because the primary paired subset is small and automatic AU proxies require manual validation.
- Automatic events and sniff-like windows are described as candidate markers requiring manual confirmation.
"""
    (ARTICLE_DIR / "README_article_sources.md").write_text(text, encoding="utf-8")


def validate_outputs(metrics: dict[str, object]) -> None:
    required = [
        ARTICLE_DIR / "main.tex",
        ARTICLE_DIR / "references.bib",
        ARTICLE_DIR / "microexpressions_article_verified.docx",
        ARTICLE_DIR / "README_article_sources.md",
        *[FIGURES_DIR / source.name for source in FIGURE_SOURCES],
    ]
    missing = [path for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError(f"missing expected outputs: {missing}")

    main_tex = (ARTICLE_DIR / "main.tex").read_text(encoding="utf-8")
    readme = (ARTICLE_DIR / "README_article_sources.md").read_text(encoding="utf-8")
    checks = [
        "450,805" in main_tex,
        "3,614" in main_tex,
        "114" in main_tex,
        "sweat/control" in main_tex,
        "WhatsApp images" in main_tex,
        "trial 1 was coded as placebo/control" in main_tex,
        "trial 2 as sweat" in main_tex,
        "Processed frame rows: 450,805" in readme,
        "Primary paired videos:" in readme,
    ]
    if not all(checks):
        raise AssertionError("content validation failed")

    with zipfile.ZipFile(ARTICLE_DIR / "microexpressions_article_verified.docx") as package:
        names = set(package.namelist())
        if "word/document.xml" not in names:
            raise AssertionError("DOCX missing word/document.xml")
        doc_text = package.read("word/document.xml").decode("utf-8", errors="ignore")
        for needle in ["450,805", "3,614", "candidate", "placebo/control"]:
            if needle not in doc_text:
                raise AssertionError(f"DOCX missing expected text: {needle}")

    if metrics["processed_frames"] != 450805 or metrics["events"] != 3614 or metrics["sniffs"] != 114:
        raise AssertionError("source metrics changed unexpectedly")


def main() -> None:
    if not WORD_SOURCE.exists():
        raise FileNotFoundError(WORD_SOURCE)
    if not PPTX_SOURCE.exists():
        raise FileNotFoundError(PPTX_SOURCE)
    ARTICLE_DIR.mkdir(parents=True, exist_ok=True)
    copy_figures()
    metrics = collect_metrics()
    write_references()
    write_main_tex(metrics)
    write_docx(metrics)
    write_readme(metrics)
    validate_outputs(metrics)
    print(f"Article draft written to: {ARTICLE_DIR}")
    print(f"Processed videos: {metrics['manifest_rows']}")
    print(f"Processed frames: {metrics['processed_frames']:,}")
    print(f"Candidate events: {metrics['events']:,}")
    print(f"Candidate sniff-like windows: {metrics['sniffs']}")
    print(f"Annotation rows: {metrics['annotations_rows']}")


if __name__ == "__main__":
    main()
